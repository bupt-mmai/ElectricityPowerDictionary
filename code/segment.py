import jieba
# import textract
# import docx2txt
import docx

# TODO: use textract to transfer doc to docx
# text = textract.process(u"../data/20170209报送企协下发的变电五通一措/国家电网公司变电检修管理规定.docx")
# print(text)

# TODO: warning this method is only for docx
# f_name = "../data/corpus/test3.doc"
f_name = '/Users/zzy824/PycharmProjects/work/ElectricityPowerDictionary/data/corpus/test3.docx'
outputfile = open(r'../result/segment3.txt', 'w',encoding='utf-8')
doc = docx.Document(f_name)

for paragraph in doc.paragraphs:
    seg = jieba.cut(paragraph.text.strip(),cut_all = False)
    output =' '.join(seg)
    outputfile.write(output + '\n')

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if len(cell.text[-2:]) > 0:
                # print('c', cell.text)
                # print(type(cell.text))
                # print(len(cell.text))
                seg = jieba.cut(cell.text.strip(),cut_all = False)
                output =' '.join(seg)
                outputfile.write(output + '\n')
outputfile.close()
    


# TODO: supplyment code
# #定义一个空字符串
# final = ""
# #文件夹位置
# # filename = r"E:\Program Files\爬虫\word.txt"
# filename=r'C:\Users\liyazhou\Desktop\a.txt'
# outputfile = open(r'C:\Users\liyazhou\Desktop\d.txt', 'w',encoding='utf-8')

# import jieba
 
# with open(r'C:\Users\liyazhou\Desktop\c.txt','rb')as f:
#     for line in f:
#       seg = jieba.cut(line.strip(),cut_all = False)
#       output =' '.join(seg)# 不是逗号？
#       # s = (r'C:\Users\liyazhou\Desktop\b.txt','a+')
#       print(output)
#       

# outputfile.close()